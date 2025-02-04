import os
import time
import fitz  # PyMuPDF for PDF file handling
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from langchain_community.embeddings import OpenAIEmbeddings
from langchain_community.vectorstores import Chroma
from pathlib import Path
from docx import Document as DocxDocument  # To handle Word files
from dotenv import load_dotenv

# Load environment variables from .env file
load_dotenv()

# Set the OpenAI API key securely from environment variables
openai_api_key = os.getenv('OPENAI_API_KEY')

# Set the folder path to monitor
FOLDER_PATH = r"C:\Users\sathv\Gen-AI\python\Python\Chunking\continuous_cunking"  # Use raw string
PROCESSED_FILES = {}  # Dictionary to track already processed files with modification times

# 1. Load a large document (handling PDF, TXT, and DOCX)
def load_document(file_path):
    file_str = str(file_path)  # Convert Path object to string
    if file_str.endswith('.pdf'):
        doc = fitz.open(file_path)
        document_text = ""
        for page_num in range(doc.page_count):
            page = doc.load_page(page_num)
            document_text += page.get_text()  # Extract text from the page
        return document_text
    elif file_str.endswith('.txt'):
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()  # Read text from a .txt file
    elif file_str.endswith('.docx'):
        doc = DocxDocument(file_path)
        document_text = ""
        for para in doc.paragraphs:
            document_text += para.text + "\n"
        return document_text
    else:
        return ""  # Handle unsupported file types

# 2. Chunk the document (splitting into smaller parts while maintaining context)
def chunk_document(document, chunk_size=1000, chunk_overlap=200):
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len
    )
    chunks = text_splitter.split_text(document)
    return chunks

# 3. Generate embeddings using OpenAI's API (or another embedding model)
def generate_embeddings(chunks):
    embeddings = OpenAIEmbeddings(api_key=openai_api_key)  # Use the securely loaded API key
    chunk_embeddings = embeddings.embed_documents(chunks)
    return chunk_embeddings

# 4. Store embeddings in Chroma DB
def store_in_chroma(chunks, embeddings):
    documents = [Document(page_content=chunk) for chunk in chunks]
    vectorstore = Chroma.from_documents(documents, OpenAIEmbeddings(), persist_directory="chroma_db")
    return vectorstore

# 5. Monitor folder for new files
def monitor_folder():
    while True:
        print("\nChecking for new files...")

        # Get all files in the folder
        files_in_folder = set(Path(FOLDER_PATH).glob('*'))
        for file in files_in_folder:
            last_modified_time = file.stat().st_mtime  # Get the file's last modified time
            
            # If the file is new or has been modified since the last check, process it
            if file not in PROCESSED_FILES or PROCESSED_FILES[file] < last_modified_time:
                print(f"\nProcessing file: {file}")
                document = load_document(file)
                
                if not document:  # Skip empty or unsupported files
                    print(f"Skipped unsupported or empty file: {file}")
                    continue
                
                # Chunk the document
                chunks = chunk_document(document)

                # Print the first 2 chunks
                print("\nChunks (first 2):")
                for i, chunk in enumerate(chunks[:2]):
                    print(f"\nChunk {i+1}:\n")
                    print(chunk)

                # Generate embeddings
                embeddings = generate_embeddings(chunks)

                # Store in Chroma DB
                store_in_chroma(chunks, embeddings)

                # Update the processed files dictionary with the new modified time
                PROCESSED_FILES[file] = last_modified_time
                print(f"Document processed and stored in Chroma DB for {file}")

        print("No new files found.")
        # Wait for 1 minute before checking again
        time.sleep(60)

# Run the folder monitoring
monitor_folder()