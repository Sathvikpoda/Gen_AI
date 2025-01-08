import os
import time
import fitz  # PyMuPDF for PDF file handling
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from sentence_transformers import SentenceTransformer  # SentenceTransformers for embeddings
import faiss  # FAISS for similarity search
from pathlib import Path
from docx import Document as DocxDocument  # To handle Word files
from dotenv import load_dotenv
import openai
import numpy as np

# Load environment variables from .env file
load_dotenv()

# Set the OpenAI API key securely from environment variables
openai_api_key = os.getenv('OPENAI_API_KEY')

# Set the folder path to monitor
FOLDER_PATH = r"C:\Users\sathv\Gen-AI\python\Python\Chunking\continuous_cunking"
PROCESSED_FILES = {}  # Dictionary to track already processed files with modification times

# Load a large document (handling PDF, TXT, and DOCX)
def load_document(file_path):
    file_str = str(file_path)  # Convert Path object to string
    document_text = ""
    try:
        if file_str.endswith('.pdf'):
            doc = fitz.open(file_path)
            for page_num in range(doc.page_count):
                page = doc.load_page(page_num)
                document_text += page.get_text()  # Extract text from the page
        elif file_str.endswith('.txt'):
            with open(file_path, 'r', encoding='utf-8') as file:
                document_text = file.read()  # Read text from a .txt file
        elif file_str.endswith('.docx'):
            doc = DocxDocument(file_path)
            for para in doc.paragraphs:
                document_text += para.text + "\n"
        else:
            return ""  # Handle unsupported file types
    except Exception as e:
        print(f"Error reading file {file_path}: {e}")
        return ""
    
    if len(document_text.strip()) == 0:
        print(f"Warning: Empty or invalid content in {file_path}. Skipping...")
        return ""
    
    return document_text

# Chunk the document (splitting into smaller parts while maintaining context)
def chunk_document(document, chunk_size=1000, chunk_overlap=200):
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len
    )
    chunks = text_splitter.split_text(document)

    # Filter out empty, invalid, or meaningless chunks
    valid_chunks = [
        chunk for chunk in chunks 
        if len(chunk.strip()) > 0 and "lorem" not in chunk.lower()
    ]
    return valid_chunks

# Generate embeddings using SentenceTransformers
def generate_embeddings():
    model = SentenceTransformer('all-MiniLM-L6-v2')  # Using a pre-trained sentence transformer model
    
    # Define and return a callable embedding function that FAISS can use
    def embedding_function(chunks):
        embeddings = model.encode(chunks, convert_to_tensor=False)  # Get embeddings
        return np.array(embeddings).astype('float32')  # Convert to np.array for FAISS compatibility
    return embedding_function

# Store embeddings in FAISS
def store_in_faiss(chunks, embedding_function):
    documents = [Document(page_content=chunk) for chunk in chunks]
    
    # Generate embeddings for the document texts
    embeddings = embedding_function([doc.page_content for doc in documents])  # Pass document contents for embedding
    
    # Initialize FAISS index
    dim = embeddings.shape[1]  # Dimensionality of the embeddings
    index = faiss.IndexFlatL2(dim)  # L2 distance-based index
    index.add(embeddings)  # Add the embeddings to the index
    
    return index, documents

# Retrieve relevant documents from FAISS based on a query
def retrieve_from_faiss(query, k=5, index=None, documents=None):
    embedding_function = generate_embeddings()
    query_embedding = embedding_function([query])  # Embed the query
    
    # Perform the search
    D, I = index.search(query_embedding, k)  # D: distances, I: indices
    
    # Retrieve corresponding documents based on the indices
    retrieved_documents = [documents[i] for i in I[0]]
    return retrieved_documents

# Agent 1: Generate response based on retrieved chunks
def generate_response_agent1(retrieved_chunks, user_query):
    context = " ".join([doc.page_content for doc in retrieved_chunks])  # Combine all retrieved chunks
    prompt = f"Extract the answer to the following query from the context below:\n\nQuery: {user_query}\nContext: {context}"

    try:
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",  # Or "gpt-4" depending on your need
            messages=[
                {"role": "system", "content": "You are a helpful assistant."},
                {"role": "user", "content": prompt},
            ],
            max_tokens=50,  # Adjust token length for short answers
            api_key=openai_api_key
        )

        return response['choices'][0]['message']['content'].strip()

    except Exception as e:
        print(f"Error generating response: {e}")
        return "Sorry, there was an error generating a response."

# Agent 2: Refine the response and make it more structured
def generate_response_agent2(agent1_output):
    # Process the raw output from Agent 1, making it more structured.
    prompt = f"Refine and structure the following information into bullet points:\n\n{agent1_output}"

    try:
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",  # Or "gpt-4" depending on your need
            messages=[
                {"role": "system", "content": "You are a helpful assistant."},
                {"role": "user", "content": prompt},
            ],
            max_tokens=100,
            api_key=openai_api_key
        )

        return response['choices'][0]['message']['content'].strip()

    except Exception as e:
        print(f"Error refining response: {e}")
        return "Sorry, there was an error refining the response."


# Monitor folder for new files and process them
def monitor_folder():
    index = None
    documents = []
    
    while True:
        print("\nChecking for new files...")

        # Get all files in the folder
        files_in_folder = set(Path(FOLDER_PATH).glob('*'))
        for file in files_in_folder:
            last_modified_time = file.stat().st_mtime  # Get the file's last modified time
            
            # If the file is new or has been modified since the last check, process it
            if file not in PROCESSED_FILES or PROCESSED_FILES[file] < last_modified_time:
                print(f"\nProcessing file: {file}")
                document = load_document(file)
                
                if not document:  # Skip empty or unsupported files
                    print(f"Skipped unsupported or empty file: {file}")
                    continue
                
                # Chunk the document
                chunks = chunk_document(document)

                # Generate embeddings using SentenceTransformers
                embeddings_function = generate_embeddings()  # Get the embedding function
                index, documents = store_in_faiss(chunks, embeddings_function)  # Store embeddings in FAISS

                # Update the processed files dictionary with the new modified time
                PROCESSED_FILES[file] = last_modified_time
                print(f"Document processed and stored in FAISS for {file}")

                # Take user query as input
                user_query = input("Enter your query: ")

                # Retrieve from FAISS and generate a response
                retrieved_chunks = retrieve_from_faiss(user_query, k=5, index=index, documents=documents)
                response_agent1 = generate_response_agent1(retrieved_chunks, user_query)  # Agent 1 response
                response_agent2 = generate_response_agent2(response_agent1)  # Agent 2 processes the output from Agent 1

                print(f"Refined Generated Response:\n{response_agent2}")

        print("No new files found.")
        # Wait for 1 minute before checking again
        time.sleep(60)

# Run the folder monitoring
monitor_folder()
